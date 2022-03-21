"""
홍보용 팩스 표지 작성
2020.11.02
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
import datetime
import pandas as pd


def word_form(d_day, company, tel, fax):
    doc = Document('D:\\2021\\_Py\\CompletionCode\\워드팩스\\홍보팩스_표지.docx')
    table = doc.tables[0]
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text == "{d_day}":
                    para.text = ""
                    p = para.add_run(" " + d_day)
                    p.font.name = "LG스마트체2.0 Regular"
                    p.element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'LG스마트체2.0 Regular')
                    p.font.size = Pt(10)
                elif para.text == "{company}":
                    para.text = ""
                    p = para.add_run(company)
                    p.font.name = "LG스마트체2.0 Regular"
                    p.element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'LG스마트체2.0 Regular')
                    p.font.size = Pt(11)
                    p.bold = True
                elif para.text == "{tel}":
                    para.text = ""
                    p = para.add_run(tel)
                    p.font.name = "LG스마트체2.0 Regular"
                    p.element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'LG스마트체2.0 Regular')
                    p.font.size = Pt(10)
                elif para.text == "{fax}":
                    para.text = ""
                    p = para.add_run(fax)
                    p.font.name = "LG스마트체2.0 Regular"
                    p.element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'LG스마트체2.0 Regular')
                    p.font.size = Pt(10)
    doc.save(f'D:\\2021\\_Py\\홍보팩스_{company}.docx')
    return None


df_temp = pd.read_excel('D:\\2021\\_Py\\CompletionCode\\워드팩스\\홍보팩스_협회목록.xlsx',
                        sheet_name='한국건축물유지관리협회',
                        header=0,
                        usecols='B,K,J')

# 발신일
now = datetime.datetime.now()
weekday_list = ['일', '월', '화', '수', '목', '금', '토']
weekday = weekday_list[int(now.strftime('%w'))]
d_day = now.strftime('%Y년 %m월 %d일 ') + weekday + '요일'

# 수신자, TEL, Fax
df_list = [2, 6, 24, 31, 32, 48, 61, 70]
for df_index in df_list:
    company = df_temp.loc[df_index-1, '업체명']
    tel = df_temp.loc[df_index-1, '연락처(Tel)']
    fax = df_temp.loc[df_index-1, '연락처(Fax)']
    word_form(d_day, company, tel, fax)

# END
